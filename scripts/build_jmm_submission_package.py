#!/usr/bin/env python3
"""Build the non-submitting upload package for Journal of Molecular Modeling."""
from __future__ import annotations

import shutil
import subprocess
import sys
import zipfile
import os
from pathlib import Path

POPPLER_BIN = Path(sys.executable).resolve().parents[1] / "native" / "poppler" / "Library" / "bin"
if POPPLER_BIN.exists():
    os.environ["PATH"] = str(POPPLER_BIN) + os.pathsep + os.environ.get("PATH", "")

import cairosvg
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "manuscript" / "jmm-submission-package"
MANUSCRIPT = ROOT / "manuscript" / "JMM-MANUSCRIPT-SUBMISSION-CANDIDATE.docx"
FIGURES = {
    "Fig1": ROOT / "reports" / "figures" / "figure-6-preparation-comparator.svg",
    "Fig2": ROOT / "reports" / "figures" / "figure-7-reference-pose-panel.svg",
}


def write(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8", newline="\n")


def normalize_eps(path: Path) -> None:
    content = path.read_text(encoding="latin-1")
    path.write_text("\n".join(line.rstrip() for line in content.splitlines()) + "\n", encoding="latin-1", newline="\n")


def build_supplement() -> None:
    archive = PACKAGE / "OnlineResource1_ReproducibilityArtefacts.zip"
    files = [
        "data/dude_targets.csv",
        "results/dude_receptor_audit.csv",
        "results/rcsb_mmcif_preparation_audit.csv",
        "results/braf_sm5_reference_pose_recovery.csv",
        "results/kif11_k30_reference_pose_recovery.csv",
        "validation/rcsb_legacy_pdb_download_manifest_20260809.csv",
        "validation/rcsb_legacy_pdb_preparation_audit_20260809.csv",
        "scripts/verify_frozen_evidence.py",
        "scripts/render_preparation_comparator_figure.py",
        "scripts/render_reference_pose_panel.py",
        "protocol/braf_3d4q_vina_reference.conf",
        "protocol/kif11_3cjo_vina_reference.conf",
        "validation/README.md",
    ]
    readme = """Supplementary information for Journal of Molecular Modeling

Article title: A provenance-aware workflow for strict receptor-preparation audits and reference-pose recovery in molecular docking
Authors: Andrés Monreal Hernández; Sara Lizbeth Franco Amaya; Carlos Ivanhoe Martínez Osorio
Affiliation: Universidad Estatal de Sonora / Universidad de Sonora, Hermosillo, Sonora, Mexico
Corresponding author: andres.monreal@ues.mx

Contents: non-coordinate derived tables, protocols, deterministic rendering scripts, and verification material supporting the manuscript.
Excluded material: third-party coordinate files and local PDBQT/log outputs are not redistributed. Source identifiers, URLs, and checksums are retained in the included records and public repository.

Use: extract this archive and consult the repository README and validation documentation. The archive is submitted for peer-review support; its versioned public archive DOI must be inserted into the manuscript before final submission.
"""
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        bundle.writestr("README.txt", readme)
        for relative in files:
            source = ROOT / relative
            if not source.exists():
                raise FileNotFoundError(source)
            bundle.write(source, relative)


def insert_picture_after(paragraph: Paragraph, image: Path) -> None:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    image_paragraph = Paragraph(node, paragraph._parent)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image), width=Inches(6.5))
    image_paragraph.paragraph_format.space_after = 0


def embed_figures() -> None:
    pngs: dict[str, Path] = {}
    for label, source in FIGURES.items():
        png = PACKAGE / f"{label}-embedded.png"
        cairosvg.svg2png(url=str(source), write_to=str(png), output_width=2400)
        pngs[label] = png
    document = Document(MANUSCRIPT)
    captions = {"Fig. 1": pngs["Fig1"], "Fig. 2": pngs["Fig2"]}
    for paragraph in list(document.paragraphs):
        for prefix, image in captions.items():
            if paragraph.text.startswith(prefix):
                insert_picture_after(paragraph, image)
    document.save(MANUSCRIPT)


def main() -> None:
    subprocess.run([sys.executable, "scripts/render_preparation_comparator_figure.py"], cwd=ROOT, check=True)
    subprocess.run([sys.executable, "scripts/render_reference_pose_panel.py"], cwd=ROOT, check=True)
    subprocess.run(
        [sys.executable, "scripts/build_jmm_author_review_docx.py", "--output", str(MANUSCRIPT)],
        cwd=ROOT,
        check=True,
    )

    PACKAGE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MANUSCRIPT, PACKAGE / "JMM-Manuscript.docx")
    for label, source in FIGURES.items():
        cairosvg.svg2pdf(url=str(source), write_to=str(PACKAGE / f"{label}.pdf"))
        cairosvg.svg2ps(url=str(source), write_to=str(PACKAGE / f"{label}.eps"))
        normalize_eps(PACKAGE / f"{label}.eps")
        shutil.copy2(source, PACKAGE / f"{label}-source.svg")
    embed_figures()
    build_supplement()

    write(PACKAGE / "COVER-LETTER-DRAFT.txt", """Dear Editor,

Please consider our manuscript, “A provenance-aware workflow for strict receptor-preparation audits and reference-pose recovery in molecular docking,” for publication as a Software Report in Journal of Molecular Modeling.

The manuscript reports a provenance-aware, strict receptor-preparation workflow for molecular docking. It preserves and classifies failures rather than silently repairing inputs, compares direct preparation outcomes across original DUD-E PDB, RCSB mmCIF, and RCSB legacy-PDB conditions for 102 entries, and documents two bounded reference-pose recovery cases with retained configurations and all generated poses. The work is intended as reproducibility-oriented computational infrastructure; it does not claim biological activity, therapeutic effect, or experimental validation.

The accompanying files provide editable manuscript source, vector figure derivatives, and a supplementary archive containing non-coordinate derived data, protocols, and verification scripts. A versioned public archive for the comparator and manuscript-source state will be created and its persistent link inserted before submission.

Before pressing Submit, the corresponding author must confirm in the portal that the manuscript is original, is not under consideration elsewhere, and has been approved by all authors and relevant institutions.

Sincerely,
Andrés Monreal Hernández
Corresponding author
andres.monreal@ues.mx
""")
    write(PACKAGE / "SUBMISSION-METADATA.txt", """ARTICLE TYPE
Software Report

TITLE
A provenance-aware workflow for strict receptor-preparation audits and reference-pose recovery in molecular docking

RUNNING TITLE
Provenance-aware receptor-preparation audits for docking

KEYWORDS
molecular docking; receptor preparation; reproducibility; provenance; Meeko; AutoDock Vina

AUTHORS
Andrés Monreal Hernández | Universidad Estatal de Sonora, Hermosillo, Sonora, Mexico | ORCID 0009-0009-1207-8597 | corresponding author | andres.monreal@ues.mx
Sara Lizbeth Franco Amaya | Doctorado en Nanotecnología, Universidad de Sonora, Hermosillo, Sonora, Mexico | ORCID 0009-0005-0272-0241
Carlos Ivanhoe Martínez Osorio | Doctorado en Ciencia de Materiales, Universidad de Sonora, Hermosillo, Sonora, Mexico | ORCID 0009-0003-7872-4965 | a223230105@unison.mx

PUBLIC REPOSITORY
https://github.com/sircalch/dude-receptor-prep-audit

FROZEN BASELINE ARCHIVE
https://doi.org/10.5281/zenodo.21866318

REQUIRED UPDATE BEFORE SUBMISSION
Create the new versioned archive containing comparator and manuscript-source material, then replace the manuscript placeholder sentence with its persistent DOI or URL.
""")
    shutil.copy2(MANUSCRIPT, PACKAGE / "JMM-Manuscript.docx")
    write(PACKAGE / "UPLOAD-ORDER.txt", """1. Upload JMM-Manuscript.docx as the editable manuscript text; it already contains Figs. 1 and 2 in the body.
2. Retain Fig1.eps and Fig2.eps as the preferred vector figure files. Upload them separately only if the portal requests figure files or the manuscript file is too large.
3. Upload OnlineResource1_ReproducibilityArtefacts.zip only if supplementary information is selected in the portal.
4. Paste or adapt COVER-LETTER-DRAFT.txt if the portal requests a cover letter.
5. Enter the title, author names, affiliations, ORCIDs, keywords, funding, competing-interest statement, author contributions, and data-availability statement from SUBMISSION-METADATA.txt.
6. Before final submission, create the new archive, update the DOI/URL in the manuscript, rebuild this package, and verify the generated portal PDF.
""")
    print(PACKAGE)


if __name__ == "__main__":
    main()
