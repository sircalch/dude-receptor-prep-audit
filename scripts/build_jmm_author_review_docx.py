#!/usr/bin/env python3
"""Build a JMM author-review DOCX from the evidence-bounded Markdown draft."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "JMM-MANUSCRIPT-AUTHOR-REVIEW-DRAFT.md"
OUTPUT = ROOT / "manuscript" / "JMM-MANUSCRIPT-AUTHOR-REVIEW-DRAFT.docx"


def set_font(run, name="Times New Roman", size=10, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def clean(text):
    text = re.sub(r"<sup>(.*?)</sup>", r"\1", text)
    return text.replace("**", "").replace("`", "")


def add_table(doc, cells):
    table = doc.add_table(rows=len(cells), cols=len(cells[0]))
    table.style = "Table Grid"
    widths = [round(9360 / len(cells[0]))] * len(cells[0])
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for row_index, values in enumerate(cells):
        for cell, value in zip(table.rows[row_index].cells, values):
            cell.text = ""
            run = cell.paragraphs[0].add_run(clean(value))
            set_font(run, size=8.5, bold=(row_index == 0))
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1
    for style_name, size in (("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    first_title = True
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
            records = [line]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                records.append(lines[index])
                index += 1
            cells = [[part.strip() for part in row.strip("|").split("|")] for row in records]
            add_table(doc, cells)
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(clean(line[2:]))
            set_font(run, size=14, bold=True)
            paragraph.paragraph_format.space_after = Pt(12)
            first_title = False
        elif line.startswith("## "):
            doc.add_paragraph(clean(line[3:]), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(clean(line[4:]), style="Heading 2")
        elif line.startswith("**") and line.endswith("**"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(clean(line))
            set_font(run, size=10, bold=True)
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(clean(line))
            set_font(run, size=10)
        index += 1

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    run = note.add_run("Figure files for separate upload: reports/figures/figure-6-preparation-comparator.svg and reports/figures/figure-5-braf-sm5-pose-recovery.svg.")
    set_font(run, size=9)
    doc.core_properties.title = "A provenance-aware workflow for strict receptor-preparation audits and reference-pose recovery in molecular docking"
    doc.core_properties.author = "Andrés Monreal Hernández; Sara Lizbeth Franco Amaya; Carlos Ivanhoe Martínez Osorio"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
